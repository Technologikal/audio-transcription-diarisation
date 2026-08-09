#!/usr/bin/env python3
"""
Build branded Heeley Trust meeting minutes .docx for the
Sum Studios Facilities Management visit — 11 March 2026.

Run:
    python3 build_minutes_docx.py
Output:
    /mnt/Technologikal/Recordings/HT/HT-Buildings-meeting-110326-minutes.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand ─────────────────────────────────────────────────────────────────────
ORANGE   = RGBColor(0xE6, 0x7A, 0x14)
DARK     = RGBColor(0x12, 0x12, 0x12)
TEXT     = RGBColor(0x27, 0x26, 0x26)
GREY     = RGBColor(0x88, 0x88, 0x88)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
L_ORANGE = "FAE5CC"   # pale orange for table headers
D_ORANGE = "E67A14"   # hex string for shading

LOGO = Path("/mnt/Technologikal/Recordings/HT/heeley_trust_logo.png")
OUT  = Path("/mnt/Technologikal/Recordings/HT/HT-Buildings-meeting-110326-minutes.docx")


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _shd(element, fill: str):
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  fill.lstrip("#"))
    element.append(s)


def set_cell_bg(cell, hex_colour: str):
    tcPr = cell._tc.get_or_add_tcPr()
    _shd(tcPr, hex_colour)


def set_para_shd(para, hex_colour: str):
    pPr = para._p.get_or_add_pPr()
    _shd(pPr, hex_colour)


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                b.append(el)
            tcPr.append(b)


def run(para, text, bold=False, italic=False, colour=None, size=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Arial"
    if colour: r.font.color.rgb = colour
    if size:   r.font.size = Pt(size)
    return r


def spacer(doc, before=0, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p


# ── Structural helpers ────────────────────────────────────────────────────────

def section_heading(doc, title: str):
    """Orange-background section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(4)
    set_para_shd(p, D_ORANGE)
    run(p, title.upper(), bold=True, colour=WHITE, size=10)
    return p


def sub_heading(doc, title: str):
    """Orange underlined sub-heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run(p, title, bold=True, colour=ORANGE, size=10)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), D_ORANGE)
    bdr.append(bot)
    pPr.append(bdr)
    return p


def bullet(doc, text: str, indent_cm=1.0, size=9):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run(p, text, colour=TEXT, size=size)
    return p


def sub_bullet(doc, text: str):
    return bullet(doc, text, indent_cm=1.8, size=9)


def body(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run(p, text, colour=TEXT, size=9)
    return p


# ── Header ────────────────────────────────────────────────────────────────────

def build_header(doc):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(t)

    # Logo cell
    lc = t.cell(0, 0)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)
    if LOGO.exists():
        lp.add_run().add_picture(str(LOGO), width=Cm(4.5))
    else:
        run(lp, "Heeley Trust", bold=True, size=16, colour=DARK)

    # Title cell
    tc = t.cell(0, 1)
    set_cell_bg(tc, D_ORANGE)

    tp = tc.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(0)
    run(tp, "MEETING MINUTES", bold=True, colour=WHITE, size=14)

    sp = tc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after  = Pt(6)
    run(sp, "Sum Studios — Facilities Management Visit", colour=WHITE, size=10)

    doc.add_paragraph()  # gap


# ── Details table ─────────────────────────────────────────────────────────────

def build_details(doc):
    rows = [
        ("Date",       "11 March 2026"),
        ("Location",   "Sum Studios, Heeley, Sheffield"),
        ("Attendees",  "Sam Reavey — Community Buildings Manager, Heeley Trust\n"
                       "Tony Askins — Lead Trustee (Buildings Sub-group), Heeley Trust\n"
                       "Anthony Ashton — Trustee (IT Specialist), Heeley Trust"),
        ("Apologies",  "None"),
        ("Minutes by", "Anthony Ashton"),
    ]

    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(rows):
        lc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        lc.width = Cm(3.2)
        vc.width = Cm(13.2)

        set_cell_bg(lc, L_ORANGE)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after  = Pt(2)
        run(lp, label, bold=True, size=9, colour=TEXT)

        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after  = Pt(2)
        run(vp, value, size=9, colour=TEXT)

    doc.add_paragraph()


# ── Purpose ───────────────────────────────────────────────────────────────────

def build_purpose(doc):
    section_heading(doc, "1.  Purpose of Meeting")
    body(doc,
         "Tony Askins opened the meeting, explaining the purpose was to "
         "build a clearer picture of Sum Studios' facilities condition, long-term "
         "maintenance costs, and capital planning requirements, so that the Buildings "
         "Sub-group can provide informed reporting to the full Heeley Trust board. "
         "The meeting included a walkthrough context of each building in the Trust's portfolio.")


# ── Sum Studios ───────────────────────────────────────────────────────────────

def build_sum_studios(doc):
    section_heading(doc, "2.  Sum Studios — Building Condition")

    # 2.1
    sub_heading(doc, "2.1  Windows (Priority Issue)")
    bullet(doc,
           "The main south-facing elevation windows are the most urgent maintenance requirement. "
           "The building is Grade II Listed, which mandates softwood frames.")
    bullet(doc,
           "A recent window replacement project was completed (after multiple setbacks including "
           "wrong measurements and remanufacturing delays across five installation visits) but "
           "the new windows continue to leak in heavy rain.")
    bullet(doc,
           "All windows on the main elevation face up the valley and are therefore exposed to "
           "the worst weather. Many have wood decay visible.")
    bullet(doc,
           "Indicative total cost for both sides of the building: ~£50,000. "
           "The two most recently replaced windows on the main elevation cost approximately £7,000.")
    bullet(doc,
           "Scaffolding has been in place for approximately two years. "
           "Sam Reavey noted that once scaffolding is erected for any section it is cost-effective "
           "to progress additional works concurrently.")
    bullet(doc,
           "Sam Reavey agreed to provide a full cost estimate and recommended phasing plan "
           "for all window replacement works across both elevations (see Action 1).",
           size=9)

    # 2.2
    sub_heading(doc, "2.2  Lighting")
    bullet(doc,
           "Fluorescent tube lighting throughout the building (~4 double-tube fittings per office). "
           "The building is approximately 13 years old; fluorescent tubes have a typical lifespan "
           "of 10–12 years and are likely approaching end-of-life.")
    bullet(doc,
           "Full LED replacement cost estimated at ~£2,000 (fittings and labour). "
           "This is not immediately urgent but should be factored into the 3–5 year capital plan.")

    # 2.3
    sub_heading(doc, "2.3  Boilers and Mechanical Plant")
    bullet(doc, "Boilers and mechanical plant are up to date and in good working order. No immediate concerns.")

    # 2.4
    sub_heading(doc, "2.4  Communal Decoration")
    bullet(doc,
           "Cyclical redecoration of communal areas is currently under way. "
           "Estimated cost: ~£5,000–£6,000 (based on previous cycle, plus ~5% annual uplift). "
           "Given the volume of footfall and the size of the building, redecoration is essentially "
           "a continuous rolling programme.")

    # 2.5
    sub_heading(doc, "2.5  Heritage Lottery Fund — Hartley Building")
    bullet(doc,
           "Heeley Trust hosted a National Architecture Heritage Fund (NAHF) event at Sum Studios "
           "the week prior to this meeting.")
    bullet(doc,
           "This has opened a renewed dialogue with the Heritage Lottery Fund regarding the Hartley "
           "building. The relevant HLF contact is keen to facilitate a new application.")
    bullet(doc,
           "The previous application was for approximately £2.5 million. "
           "Andy Mears is progressing this and the full board should expect a positive update "
           "at the next board meeting. (Note: treat as confidential until Andy formally reports.)")

    # 2.6
    sub_heading(doc, "2.6  Tenancy and Rental")
    bullet(doc,
           "Rental structure: £9 per sq ft base rent plus a service charge (shared building "
           "running costs, split proportionally across tenants). Considered competitive with "
           "Sheffield city centre alternatives.")
    bullet(doc,
           "One studio was recently let; a further viewing was booked for the following day. "
           "Sam Reavey is preparing service charge increase letters to all tenants (7% increase — "
           "see Action 3).")
    bullet(doc,
           "The largest tenant (a computer games company paying ~£3,000/month) has historically "
           "had irregular payment patterns due to contract-based income. Flagged as a tenancy risk "
           "to monitor.")


# ── Ashtree Yard ──────────────────────────────────────────────────────────────

def build_ashtree(doc):
    section_heading(doc, "3.  Ashtree Yard")

    sub_heading(doc, "3.1  Staff Consolidation")
    bullet(doc,
           "Sam Reavey, Simon and Katya have all relocated from Ashtree Yard to Sum Studios. "
           "Andy Mears continues to divide his time between sites. "
           "The vacated offices at Ashtree Yard are being cleared to generate additional rental income.")
    bullet(doc,
           "Note: Sam's predecessor (Matthew) had a studio at Sum Studios as part of his "
           "remuneration package. No free studio is included in the current arrangements.")

    sub_heading(doc, "3.2  Festival Designer — Ruthie")
    bullet(doc,
           "Ruthie, a festival set designer, is taking over the top floor space at Ashtree Yard "
           "(currently being cleared). Her rental contribution will double, which will fund "
           "insulation improvements to the top office space.")
    bullet(doc,
           "The top office is currently unfit for purpose in winter (cold, uninsulated). "
           "Sam Reavey is targeting insulation works to be costed and completed by summer 2026 "
           "(see Action 4).")

    sub_heading(doc, "3.3  Ground Floor Storage — Damp Remediation")
    bullet(doc,
           "The ground floor right-hand side (currently used for bike/storage) is damp and "
           "unfit for purpose. Tanking and waterproofing works are required.")
    bullet(doc,
           "This is a medium-to-long-term project; income from the Ruthie tenancy is intended "
           "to contribute to funding it over time.")

    sub_heading(doc, "3.4  HEC (Ecological Surveying Company)")
    bullet(doc,
           "HEC, connected to the University of Sheffield, occupies the left-hand ground floor office "
           "and has been a stable tenant for approximately 5–6 years. No changes.")


# ── Harley ────────────────────────────────────────────────────────────────────

def build_harley(doc):
    section_heading(doc, "4.  Harley Building")
    bullet(doc,
           "An artist (long-standing connection to the Trust; contributed to the original bike shop "
           "design) occupies the Harley Building at peppercorn rent. He uses it as an art studio "
           "for approximately one month per year.")
    bullet(doc,
           "This arrangement is deliberately maintained: an occupied building attracts lower "
           "insurance premiums than a fully vacant property. The artist also undertakes minor "
           "maintenance and general upkeep.")
    bullet(doc,
           "Gleadless also remains on the empty buildings insurance premium at an additional "
           "~£1,000/year above a normally occupied property. No immediate change planned.")


# ── Bike Shop ─────────────────────────────────────────────────────────────────

def build_bike_shop(doc):
    section_heading(doc, "5.  Heeley Green — Bike Shop")
    bullet(doc,
           "Relatively new building with PVC windows and low ongoing maintenance requirements. "
           "Shutters serviced annually; windows cleaned quarterly; gutters cleared regularly.")
    bullet(doc,
           "Ange to remain based at Heeley Green. Future of the bike operation under review "
           "following staffing restructure, but the building itself presents no material concerns.")


# ── Library / Terry Wright ────────────────────────────────────────────────────

def build_library(doc):
    section_heading(doc, "6.  Gleadless Library / Terry Wright Building")

    sub_heading(doc, "6.1  Handover and Condition")
    bullet(doc,
           "Heeley Trust has formally taken over this site from Sheffield City Council. "
           "Key handover is scheduled for Monday 16 March 2026.")
    bullet(doc,
           "The council committed to replacing the windows prior to handover but measured "
           "incorrectly; replacement windows are on order. The existing windows are in poor "
           "condition (some frames rotten). Building survey has been completed.")
    bullet(doc,
           "The library building itself has received minimal investment; the Terry Wright "
           "building and the new extension are in better condition following recent refurbishment.")

    sub_heading(doc, "6.2  NHS Mental Health Pilot (Year 1)")
    bullet(doc,
           "An NHS-funded mental health pilot will operate from the Terry Wright building "
           "and new extension. The pilot runs for one year.")
    bullet(doc, "Facilities breakdown:")
    sub_bullet(doc, "Terry Wright building: crisis café (open 8am–10pm, drop-in), "
                    "hireable community space, and NHS/HT staff consultation rooms.")
    sub_bullet(doc, "New extension: three self-check-in beds for people in acute mental health crisis "
                    "(managed entirely by the NHS — not Heeley Trust).")
    bullet(doc,
           "Heeley Trust will be the landlord for the whole site under a full repairing lease "
           "from Sheffield City Council. A walkaway clause is in place if the pilot ends after "
           "Year 1. (Sam Reavey/Andy Mears to confirm lease terms before handover — see Action 5.)")

    sub_heading(doc, "6.3  Library Operation")
    bullet(doc,
           "The library continues to operate as a community resource (book loans, internet access). "
           "Predominantly volunteer-run; Jim (paid Volunteer Manager) coordinates volunteers.")
    bullet(doc,
           "Sheffield City Council provides a grant contribution towards library running costs. "
           "Gleadless is a deprived area; the library provides critical free internet and skills "
           "access to the local community.")

    sub_heading(doc, "6.4  Risks Noted")
    bullet(doc,
           "Antisocial behaviour and vandalism are prevalent in the area — flagged as an "
           "operational risk for the site.")
    bullet(doc,
           "Clarity on NHS's ongoing commitment and funding beyond Year 1 is still to be "
           "confirmed. The group noted the project planning has involved significant uncertainty.")


# ── IT ────────────────────────────────────────────────────────────────────────

def build_it(doc):
    section_heading(doc, "7.  IT Infrastructure Review")

    sub_heading(doc, "7.1  Current Setup")
    bullet(doc, "IT support: Clinix Group (remote support ticketing, antivirus, infrastructure management).")
    bullet(doc,
           "File sharing/internal comms: Microsoft 365 / SharePoint (cloud, monthly per-user subscription).")
    bullet(doc,
           "Phone system: 3CX (cloud-hosted VoIP). One inbound line split to individual tenant "
           "DDIs. Virgin Media broadband at Sum Studios; City Fibre at Ashtree Yard.")
    bullet(doc,
           "Websites: three separate WordPress sites (Sum Studios, Heeley Community Team, Bike Shop) "
           "plus a trust-wide site. Backend hosting and updates managed by ASG Media on a monthly "
           "retainer. Staff manage content internally.")

    sub_heading(doc, "7.2  Recent Changes")
    bullet(doc,
           "UPS (server/comms rack) replaced in late 2025 following failure — cost ~£2,500.")
    bullet(doc,
           "New laptops purchased for all staff in 2025, triggered by Windows 11 hardware "
           "requirements rendering several existing devices obsolete.")
    bullet(doc,
           "All PSTN copper lines decommissioned (BT switch-off). Alarm systems converted to "
           "wireless — generating a cost saving of several hundred pounds per year on line rentals.")

    sub_heading(doc, "7.3  Forward Planning")
    bullet(doc,
           "Laptop replacement cycle: Anthony Ashton recommended planning for replacement on a "
           "3–5 year basis. Current laptops purchased 2025; budget provision needed from ~2028.")
    bullet(doc,
           "Network switches: the switches running the building infrastructure may be approaching "
           "end of life. Anthony Ashton to review ages and provide a replacement forecast for "
           "inclusion in the capital plan (see Action 6).")
    bullet(doc,
           "Website rationalisation: Sam Reavey noted a desire to consolidate the multiple "
           "separate websites to improve clarity for the public. Flagged as a future project.")
    bullet(doc,
           "Broadband contracts: Virgin (Sum Studios) contract due for renewal — worth reviewing "
           "at renewal to compare City Fibre rates.")


# ── Actions ───────────────────────────────────────────────────────────────────

def build_actions(doc):
    section_heading(doc, "8.  Action Items")

    actions = [
        ("1", "Provide full cost estimate and recommended phasing plan for window replacement "
              "across all elevations of Sum Studios.",
         "Sam Reavey", "Before next board meeting"),
        ("2", "Progress Heritage Lottery Fund dialogue — write to HLF contacts to facilitate "
              "a new application for the Hartley building.",
         "Andy Mears", "Ongoing"),
        ("3", "Issue 7% service charge increase letters to all Sum Studios tenants.",
         "Sam Reavey", "Near-term"),
        ("4", "Obtain costings for insulation upgrade works to Ashtree Yard top office. "
              "Commission works once costed.",
         "Sam Reavey", "Summer 2026"),
        ("5", "Confirm lease terms for Gleadless Library / Terry Wright site — specifically "
              "the walkaway clause and full repairing obligations — before key handover.",
         "Sam Reavey / Andy Mears", "Before 16 March 2026"),
        ("6", "Review ages and specifications of network switches at Sum Studios and Ashtree Yard; "
              "provide an end-of-life forecast for inclusion in the capital equipment budget.",
         "Anthony Ashton", "Next meeting"),
        ("7", "Arrange next meeting at Terry Wright building (Gleadless Library site).",
         "Tony Askins", "~May 2026"),
    ]

    t = doc.add_table(rows=1 + len(actions), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    col_widths = [Cm(0.8), Cm(9.5), Cm(3.5), Cm(2.8)]

    # Header row
    for i, hdr in enumerate(["#", "Action", "Owner", "Target"]):
        c = t.rows[0].cells[i]
        c.width = col_widths[i]
        set_cell_bg(c, D_ORANGE)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run(p, hdr, bold=True, colour=WHITE, size=9)

    # Data rows
    for j, (num, action, owner, target) in enumerate(actions):
        row = t.rows[j + 1]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        row.cells[2].width = col_widths[2]
        row.cells[3].width = col_widths[3]

        if j % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, "FDF3E8")   # very light orange stripe

        for ci, text in enumerate([num, action, owner, target]):
            p = row.cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run(p, text, size=9, colour=TEXT)


# ── Next meeting ──────────────────────────────────────────────────────────────

def build_next_meeting(doc):
    section_heading(doc, "9.  Next Meeting")
    body(doc,
         "The group agreed to hold the next meeting at the Terry Wright building "
         "(Gleadless Library site) to allow a walk-around of that site. "
         "Approximately two months hence (May 2026). Tony Askins to arrange.")


# ── Footer ────────────────────────────────────────────────────────────────────

def build_footer(doc):
    doc.add_paragraph()

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(6)
    pPr = div._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), D_ORANGE)
    bdr.append(bot)
    pPr.append(bdr)

    # Charity info
    fi = doc.add_paragraph()
    fi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(fi, "Heeley Trust  |  Charity No: 1067567  |  Company Number: 3288676",
        size=7, colour=GREY, italic=True)

    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(disc,
        "These minutes are a summary record of the meeting. "
        "Please advise the minutes secretary of any corrections within 14 days.",
        size=7, colour=GREY, italic=True)

    # Signature block
    doc.add_paragraph()
    sub_heading(doc, "Approved by Chair")
    st = doc.add_table(rows=2, cols=3)
    st.style = "Table Grid"
    for ci, label in enumerate(["Signature", "Name", "Date"]):
        set_cell_bg(st.rows[0].cells[ci], L_ORANGE)
        p = st.rows[0].cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run(p, label, bold=True, size=9)
    for cell in st.rows[1].cells:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(18)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    build_header(doc)
    build_details(doc)
    build_purpose(doc)
    build_sum_studios(doc)
    build_ashtree(doc)
    build_harley(doc)
    build_bike_shop(doc)
    build_library(doc)
    build_it(doc)
    build_actions(doc)
    build_next_meeting(doc)
    build_footer(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
