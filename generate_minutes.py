#!/usr/bin/env python3
"""
Generate Heeley Trust branded meeting minutes from a diarised transcript.

Usage:
    python3 generate_minutes.py <transcript_file> [options]
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Heeley Trust Brand ────────────────────────────────────────────────────────
HT_ORANGE     = RGBColor(0xE6, 0x7A, 0x14)   # #e67a14
HT_DARK       = RGBColor(0x12, 0x12, 0x12)   # #121212
HT_LIGHT_BG   = RGBColor(0xF4, 0xF5, 0xF7)   # #f4f5f7
HT_TEXT       = RGBColor(0x27, 0x26, 0x26)   # #272626
HT_MID_GREY   = RGBColor(0x88, 0x88, 0x88)   # mid grey for secondary text
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

LOGO_PATH = Path("/mnt/Technologikal/Recordings/HT/heeley_trust_logo.png")

# Known speaker mapping — update as needed after reviewing transcript
DEFAULT_SPEAKER_MAP = {
    "SPEAKER_00": "Sam Reavey",
    "SPEAKER_01": "Tony Askins",
    "SPEAKER_02": "Anthony Ashton",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_colour: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour.lstrip("#"))
    tcPr.append(shd)


def set_paragraph_border_bottom(paragraph, colour_hex: str, size_pt: int = 12):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 8))   # eighths of a point
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), colour_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def remove_paragraph_spacing(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def add_run(paragraph, text: str, bold=False, italic=False,
            colour: RGBColor = None, size_pt: int = None, font_name: str = "Arial"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.name = font_name
    return run


def set_doc_margins(doc, top=2.0, bottom=2.0, left=2.0, right=2.0):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


# ── Transcript Parsing ────────────────────────────────────────────────────────

def parse_transcript(path: Path, speaker_map: dict) -> list[dict]:
    """
    Parse a diarised transcript into a list of speaker turns.

    Expected line format (from transcribe.py):
        [HH:MM:SS - HH:MM:SS] SPEAKER_XX: text...
    or  [MM:SS - MM:SS] SPEAKER_XX: text...
    """
    pattern = re.compile(
        r"\[(\d+:\d+(?::\d+)?)\s*-\s*(\d+:\d+(?::\d+)?)\]\s*([\w_]+):\s*(.*)"
    )
    turns = []
    current = None

    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            m = pattern.match(line)
            if m:
                start, end, speaker_raw, text = m.groups()
                speaker = speaker_map.get(speaker_raw, speaker_raw)
                # Merge consecutive same-speaker lines
                if current and current["speaker"] == speaker:
                    current["text"] += " " + text.strip()
                    current["end"] = end
                else:
                    if current:
                        turns.append(current)
                    current = {
                        "start": start,
                        "end": end,
                        "speaker": speaker,
                        "text": text.strip(),
                    }
            else:
                # Continuation line (no timestamp) — append to current turn
                if current and line:
                    current["text"] += " " + line

    if current:
        turns.append(current)

    return turns


def extract_action_items(turns: list[dict]) -> list[dict]:
    """
    Heuristic scan for action items — sentences containing action keywords.
    """
    keywords = [
        r"\bwill\b", r"\bto do\b", r"\baction\b", r"\bneed to\b", r"\bshould\b",
        r"\bmust\b", r"\bfollow[- ]?up\b", r"\bsend\b", r"\bcheck\b",
        r"\barrange\b", r"\bbook\b", r"\bcontact\b", r"\bprocure\b",
        r"\bget quotes?\b", r"\borganis[e|ing]\b", r"\bschedule\b",
    ]
    pattern = re.compile("|".join(keywords), re.IGNORECASE)
    items = []
    for turn in turns:
        for sentence in re.split(r"(?<=[.!?])\s+", turn["text"]):
            if pattern.search(sentence) and len(sentence) > 20:
                items.append({
                    "speaker": turn["speaker"],
                    "timestamp": turn["start"],
                    "text": sentence.strip(),
                })
    return items


# ── Document Building ─────────────────────────────────────────────────────────

def build_header(doc: Document, meeting_meta: dict):
    """Build the branded header: logo left, title area right."""
    # Header table: logo | meeting title block
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Remove all borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)

    # Logo cell
    logo_cell = table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        run = logo_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(5))
    else:
        add_run(logo_para, "Heeley Trust", bold=True, size_pt=16, colour=HT_DARK)

    # Title cell
    title_cell = table.cell(0, 1)
    title_cell.width = Cm(12)

    # Orange background for title cell
    set_cell_background(title_cell, "e67a14")

    # Meeting title
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(2)
    add_run(tp, "MEETING MINUTES", bold=True, size_pt=16, colour=WHITE)

    # Sub-title
    sp = title_cell.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)
    add_run(sp, meeting_meta.get("title", ""), bold=False, size_pt=11, colour=WHITE)

    doc.add_paragraph()  # spacer


def build_details_table(doc: Document, meeting_meta: dict):
    """Build the meeting details block (date, time, location, attendees)."""
    rows = [
        ("Date",      meeting_meta.get("date", "")),
        ("Time",      meeting_meta.get("time", "")),
        ("Location",  meeting_meta.get("location", "")),
        ("Attendees", "\n".join(meeting_meta.get("attendees", []))),
        ("Apologies", meeting_meta.get("apologies", "None")),
        ("Minutes by", meeting_meta.get("minutes_by", "")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    for i, row in enumerate(table.rows):
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13)

        label, value = rows[i]

        # Label cell (light orange background)
        lc = row.cells[0]
        set_cell_background(lc, "FAE5CC")
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)
        add_run(lp, label, bold=True, size_pt=9, colour=HT_TEXT)

        # Value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after = Pt(2)
        add_run(vp, value, size_pt=9, colour=HT_TEXT)

    doc.add_paragraph()  # spacer


def add_section_heading(doc: Document, title: str, level: int = 1):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    if level == 1:
        run = add_run(p, title.upper(), bold=True, size_pt=11, colour=WHITE)
        # Orange shaded background via table trick
        # Use a single-row single-cell table for the heading background
        # (Simpler: set paragraph shading)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "e67a14")
        pPr.append(shd)
        # Add left indent for padding feel
        p.paragraph_format.left_indent = Pt(6)
    else:
        add_run(p, title, bold=True, size_pt=10, colour=HT_ORANGE)
        set_paragraph_border_bottom(p, "e67a14", size_pt=1)

    return p


def build_transcript_section(doc: Document, turns: list[dict]):
    """Add the full diarised transcript as formatted discussion record."""
    add_section_heading(doc, "Discussion Record")

    for turn in turns:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0)

        # Timestamp + Speaker name
        ts_text = f"[{turn['start']}]  "
        add_run(p, ts_text, size_pt=8, colour=HT_MID_GREY, italic=True)
        add_run(p, turn["speaker"] + ":  ", bold=True, size_pt=9, colour=HT_ORANGE)
        add_run(p, turn["text"], size_pt=9, colour=HT_TEXT)


def build_action_items_section(doc: Document, items: list[dict]):
    """Add action items table."""
    add_section_heading(doc, "Action Items")

    if not items:
        p = doc.add_paragraph()
        add_run(p, "No specific action items identified from transcript.", italic=True,
                size_pt=9, colour=HT_MID_GREY)
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, col_title in enumerate(["#", "Action", "Owner", "Timestamp"]):
        cell = hdr.cells[i]
        set_cell_background(cell, "e67a14")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, col_title, bold=True, size_pt=9, colour=WHITE)

    for idx, item in enumerate(items, 1):
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(item["text"]).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(item["speaker"]).font.size = Pt(9)
        row.cells[3].paragraphs[0].add_run(item["timestamp"]).font.size = Pt(9)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)


def build_footer(doc: Document, meeting_meta: dict):
    """Add a branded footer section."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    set_paragraph_border_bottom(p, "e67a14", size_pt=1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Heeley Trust  |  Charity No: 1067567  |  Company Number: 3288676",
            size_pt=7, colour=HT_MID_GREY, italic=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "These minutes are a summary record of the meeting. "
            "Please notify the minutes secretary of any corrections within 14 days.",
            size_pt=7, colour=HT_MID_GREY, italic=True)

    # Signature block
    doc.add_paragraph()
    add_section_heading(doc, "Approval", level=2)
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.style = "Table Grid"
    sig_hdr = sig_table.rows[0]
    for i, label in enumerate(["Signed (Chair)", "Name", "Date"]):
        set_cell_background(sig_hdr.cells[i], "FAE5CC")
        p = sig_hdr.cells[i].paragraphs[0]
        add_run(p, label, bold=True, size_pt=9)

    # Empty signature row
    sig_row = sig_table.rows[1]
    for cell in sig_row.cells:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(14)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate Heeley Trust branded meeting minutes from a transcript."
    )
    parser.add_argument("transcript", help="Path to diarised transcript file")
    parser.add_argument("-o", "--output", help="Output .docx file path")
    parser.add_argument("--title",    default="Sum Studios — Facilities Management Visit",
                        help="Meeting title")
    parser.add_argument("--date",     default="11 March 2026",   help="Meeting date")
    parser.add_argument("--time",     default="",                help="Meeting time range")
    parser.add_argument("--location", default="Sum Studios",     help="Meeting location")
    parser.add_argument("--speaker-map", nargs="*", metavar="LABEL=Name",
                        help="Speaker mappings e.g. SPEAKER_00='Sam Reavey'")
    args = parser.parse_args()

    # Build speaker map
    speaker_map = dict(DEFAULT_SPEAKER_MAP)
    if args.speaker_map:
        for item in args.speaker_map:
            if "=" in item:
                k, v = item.split("=", 1)
                speaker_map[k.strip()] = v.strip().strip("'\"")

    transcript_path = Path(args.transcript)
    if not transcript_path.exists():
        print(f"ERROR: Transcript not found: {transcript_path}", file=sys.stderr)
        sys.exit(1)

    # Parse transcript
    print(f"Parsing transcript: {transcript_path}")
    turns = parse_transcript(transcript_path, speaker_map)
    print(f"  → {len(turns)} speaker turns")

    # Extract action items
    actions = extract_action_items(turns)
    print(f"  → {len(actions)} potential action items")

    # Meeting metadata
    meeting_meta = {
        "title":    args.title,
        "date":     args.date,
        "time":     args.time,
        "location": args.location,
        "attendees": [
            "Sam Reavey — Community Buildings Manager, Heeley Trust",
            "Tony Askins — Lead Trustee (Buildings Sub-group), Heeley Trust",
            "Anthony Ashton — Trustee (IT Specialist), Heeley Trust",
        ],
        "apologies":  "None",
        "minutes_by": "Anthony Ashton",
    }

    # Build document
    doc = Document()
    set_doc_margins(doc, top=1.5, bottom=2.0, left=2.0, right=2.0)

    # Set default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    build_header(doc, meeting_meta)
    build_details_table(doc, meeting_meta)
    build_transcript_section(doc, turns)
    build_action_items_section(doc, actions)
    build_footer(doc, meeting_meta)

    # Output path
    if args.output:
        out_path = Path(args.output)
    else:
        out_path = transcript_path.with_suffix(".docx")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\nMinutes saved to: {out_path}")


if __name__ == "__main__":
    main()
