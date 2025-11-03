"""
Agenda Parser Module

Parses community meeting agenda documents (DOCX format) to extract:
- Meeting metadata (date, time, location)
- Agenda sections and topics
- Speaker assignments for each section
"""

import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from docx import Document
import logging

logger = logging.getLogger(__name__)


@dataclass
class AgendaMetadata:
    """Meeting metadata extracted from agenda"""
    date: Optional[str] = None
    time: Optional[str] = None
    location: Optional[str] = None
    title: Optional[str] = None
    raw_metadata: Dict[str, str] = field(default_factory=dict)


@dataclass
class AgendaSection:
    """Individual section or topic in the agenda"""
    section_number: Optional[str] = None
    title: str = ""
    speakers: List[str] = field(default_factory=list)
    description: str = ""
    estimated_duration: Optional[int] = None  # in minutes
    subsections: List['AgendaSection'] = field(default_factory=list)

    def get_all_speakers(self) -> List[str]:
        """Get all speakers including from subsections"""
        all_speakers = list(self.speakers)
        for subsection in self.subsections:
            all_speakers.extend(subsection.get_all_speakers())
        return list(set(all_speakers))  # Remove duplicates


@dataclass
class ParsedAgenda:
    """Complete parsed agenda structure"""
    metadata: AgendaMetadata
    sections: List[AgendaSection]
    all_speakers: List[str] = field(default_factory=list)

    def __post_init__(self):
        """Extract unique list of all speakers after initialization"""
        speakers = set()
        for section in self.sections:
            speakers.update(section.get_all_speakers())
        self.all_speakers = sorted(list(speakers))


class AgendaParser:
    """
    Parser for community meeting agenda documents.

    Handles semi-structured DOCX files with flexible patterns for:
    - Metadata extraction (date, time, location)
    - Section/topic identification
    - Speaker name extraction
    """

    # Common patterns for metadata fields
    METADATA_PATTERNS = {
        'date': [
            r'date[:\s]+([^\n]+)',
            r'meeting date[:\s]+([^\n]+)',
            r'when[:\s]+([^\n]+)',
        ],
        'time': [
            r'time[:\s]+([^\n]+)',
            r'meeting time[:\s]+([^\n]+)',
            r'(\d{1,2}:\d{2}\s*(?:am|pm|AM|PM))',
        ],
        'location': [
            r'location[:\s]+([^\n]+)',
            r'venue[:\s]+([^\n]+)',
            r'where[:\s]+([^\n]+)',
            r'place[:\s]+([^\n]+)',
        ],
        'title': [
            r'meeting[:\s]+([^\n]+)',
            r'subject[:\s]+([^\n]+)',
        ]
    }

    # Patterns for identifying speakers
    SPEAKER_PATTERNS = [
        r'(?:speaker|presenter|presented by|facilitator|chair|led by)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'(?:by|with)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'\(([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\)',
        r'-\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
    ]

    # Patterns for section numbering
    SECTION_NUMBER_PATTERNS = [
        r'^(\d+\.?\d*)\s+',  # 1. or 1.1
        r'^([A-Z]\.)\s+',     # A.
        r'^([IVX]+\.)\s+',    # Roman numerals
    ]

    def __init__(self, docx_path: str):
        """
        Initialize parser with path to DOCX agenda file.

        Args:
            docx_path: Path to the DOCX agenda document
        """
        self.docx_path = docx_path
        self.document = None

    def parse(self) -> ParsedAgenda:
        """
        Parse the agenda document and extract all information.

        Returns:
            ParsedAgenda object with metadata, sections, and speakers
        """
        logger.info(f"Parsing agenda document: {self.docx_path}")

        try:
            self.document = Document(self.docx_path)
        except Exception as e:
            logger.error(f"Failed to load DOCX file: {e}")
            raise RuntimeError(f"Failed to load agenda document: {e}")

        # Extract metadata from document
        metadata = self._extract_metadata()

        # Extract agenda sections and topics
        sections = self._extract_sections()

        parsed_agenda = ParsedAgenda(metadata=metadata, sections=sections)

        logger.info(f"Parsed agenda: {len(sections)} sections, {len(parsed_agenda.all_speakers)} unique speakers")
        logger.debug(f"Speakers found: {', '.join(parsed_agenda.all_speakers)}")

        return parsed_agenda

    def _extract_metadata(self) -> AgendaMetadata:
        """
        Extract meeting metadata from the document.

        Returns:
            AgendaMetadata object with date, time, location, etc.
        """
        metadata = AgendaMetadata()

        # Search through first 20 paragraphs for metadata (typically at top)
        search_text = "\n".join([p.text for p in self.document.paragraphs[:20]])

        for field, patterns in self.METADATA_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, search_text, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    setattr(metadata, field, value)
                    metadata.raw_metadata[field] = value
                    logger.debug(f"Found {field}: {value}")
                    break

        return metadata

    def _extract_sections(self) -> List[AgendaSection]:
        """
        Extract agenda sections, topics, and speaker assignments.

        Returns:
            List of AgendaSection objects
        """
        sections = []
        current_section = None

        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading/section title
            is_heading = self._is_heading(para)

            if is_heading:
                # Save previous section if exists
                if current_section and current_section.title:
                    sections.append(current_section)

                # Start new section
                current_section = AgendaSection()
                current_section.title = text

                # Extract section number if present
                section_num = self._extract_section_number(text)
                if section_num:
                    current_section.section_number = section_num
                    current_section.title = text[len(section_num):].strip()

                # Try to extract speakers from title
                speakers = self._extract_speakers(text)
                current_section.speakers.extend(speakers)

            elif current_section:
                # Add description to current section
                if current_section.description:
                    current_section.description += "\n" + text
                else:
                    current_section.description = text

                # Try to extract speakers from description
                speakers = self._extract_speakers(text)
                current_section.speakers.extend(speakers)

        # Add final section
        if current_section and current_section.title:
            sections.append(current_section)

        # Deduplicate speakers in each section
        for section in sections:
            section.speakers = list(set(section.speakers))

        return sections

    def _is_heading(self, paragraph) -> bool:
        """
        Determine if a paragraph is a heading/section title.

        Args:
            paragraph: python-docx paragraph object

        Returns:
            True if paragraph appears to be a heading
        """
        # Check if paragraph has a heading style
        if paragraph.style and paragraph.style.name.startswith('Heading'):
            return True

        # Check if text starts with section number
        text = paragraph.text.strip()
        for pattern in self.SECTION_NUMBER_PATTERNS:
            if re.match(pattern, text):
                return True

        # Check if paragraph is bold and short (likely a heading)
        if paragraph.runs:
            # Check if most/all runs are bold
            bold_chars = sum(len(run.text) for run in paragraph.runs if run.bold)
            total_chars = sum(len(run.text) for run in paragraph.runs)
            if total_chars > 0 and bold_chars / total_chars > 0.7:
                # Short bold text is likely a heading
                if len(text) < 100:
                    return True

        return False

    def _extract_section_number(self, text: str) -> Optional[str]:
        """
        Extract section number from text (e.g., "1.", "1.1", "A.")

        Args:
            text: Section title text

        Returns:
            Section number if found, None otherwise
        """
        for pattern in self.SECTION_NUMBER_PATTERNS:
            match = re.match(pattern, text)
            if match:
                return match.group(1)
        return None

    def _extract_speakers(self, text: str) -> List[str]:
        """
        Extract speaker names from text using various patterns.

        Args:
            text: Text to search for speaker names

        Returns:
            List of speaker names found
        """
        speakers = []

        for pattern in self.SPEAKER_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                speaker = match.group(1).strip()
                # Validate it looks like a name (at least 2 words, capitalized)
                if self._is_valid_name(speaker):
                    speakers.append(speaker)

        return speakers

    def _is_valid_name(self, name: str) -> bool:
        """
        Validate that a string looks like a person's name.

        Args:
            name: Potential name string

        Returns:
            True if string appears to be a valid name
        """
        # Must have at least first and last name
        parts = name.split()
        if len(parts) < 2:
            return False

        # Each part should start with capital letter
        if not all(part[0].isupper() for part in parts):
            return False

        # Shouldn't be too long (probably not a name if > 4 words)
        if len(parts) > 4:
            return False

        # Common false positives to exclude
        excluded_words = {'Meeting', 'Agenda', 'Community', 'Board', 'Committee', 'Trust', 'Council'}
        if any(word in parts for word in excluded_words):
            return False

        return True


def parse_agenda(docx_path: str) -> ParsedAgenda:
    """
    Convenience function to parse an agenda document.

    Args:
        docx_path: Path to DOCX agenda file

    Returns:
        ParsedAgenda object
    """
    parser = AgendaParser(docx_path)
    return parser.parse()
